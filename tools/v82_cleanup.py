from pathlib import Path
import re, json, hashlib
from docx import Document
from openpyxl import load_workbook

ROOT = Path('.')
SCHOOL = 'SMP Negeri 1 Susukan'
AUTHOR = 'Sunarso, S.Pd.I, Gr'
ALLAH = "Alloh Subhanahu Wata'ala"
MUHAMMAD = "Muhammad Sholallohu 'Alaihi Wasalam"


def std(text):
    if not isinstance(text, str):
        return text
    text = re.sub(r'SMP\s+Negeri\s+1\s+Kebonagung', SCHOOL, text, flags=re.I)
    text = re.sub(r'Syaekudin\s*,?\s*S\.?Ag\.?\s*,?\s*M\.?Pd\.?I\.?', AUTHOR, text, flags=re.I)
    text = re.sub(r'\b(?:Allah|Alloh)\s+(?:Swt|SWT)\.?', ALLAH, text, flags=re.I)
    text = re.sub(r'\bAllah\s+Subhanahu\s+Wata[’\']ala\b', ALLAH, text, flags=re.I)
    text = re.sub(r'\b(?:salat|shalat)\b', 'sholat', text, flags=re.I)
    text = re.sub(r'\bzikir\b', 'dzikir', text, flags=re.I)
    text = re.sub(r'\bhusnuzan\b', 'husnudzon', text, flags=re.I)
    text = re.sub(r'\bhadis\b', 'hadits', text, flags=re.I)
    text = re.sub(r'\bMuhammad\s+(?:saw|SAW)\.?', MUHAMMAD, text, flags=re.I)
    text = re.sub(r'Penyusun\s*:\s*' + re.escape(AUTHOR), 'Penyusun : ' + AUTHOR, text, flags=re.I)
    return text


sig_rules = [
    re.compile(r'^\s*Demak\s*,\s*Juli\s*2026\s*$', re.I),
    re.compile(r'^\s*Mengetahui\s*,?\s*$', re.I),
    re.compile(r'^\s*Kepala\s+SMP\s+Negeri\s+1\s+Kebonagung\s*,?\s*$', re.I),
    re.compile(r'Priyantono\s*,?\s*S\.?Pd\.?\s*,?\s*M\.?Pd\.?', re.I),
    re.compile(r'196902251994031005'),
    re.compile(r'^\s*Guru\s+Mata\s+Pelajaran\s+PAI(?:\s+dan\s+Budi\s+Pekerti)?\s*,?\s*$', re.I),
    re.compile(r'Syaekudin\s*,?\s*S\.?Ag\.?\s*,?\s*M\.?Pd\.?I\.?', re.I),
    re.compile(r'197209052005011004'),
    re.compile(r'^\s*NIP\.?\s*(?:196902251994031005|197209052005011004)\s*$', re.I),
]


def signature(text):
    t = ' '.join(str(text or '').split())
    return any(p.search(t) for p in sig_rules)


def remove_paragraph(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_table(table):
    el = table._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def set_paragraph_text(p, value):
    if p.text == value:
        return
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(value)


def table_text(table):
    return ' '.join(cell.text for row in table.rows for cell in row.cells)


def clean_table(table):
    original = table_text(table)
    if signature(original):
        remove_table(table)
        return False
    for row in list(table.rows):
        for cell in row.cells:
            for p in list(cell.paragraphs):
                if signature(p.text):
                    remove_paragraph(p)
                else:
                    set_paragraph_text(p, std(p.text))
            for nested in list(cell.tables):
                clean_table(nested)
        if not any(' '.join(cell.text.split()) for cell in row.cells):
            try:
                table._tbl.remove(row._tr)
            except Exception:
                pass
    if not ' '.join(table_text(table).split()):
        remove_table(table)
        return False
    return True


def clean_container(container):
    for p in list(container.paragraphs):
        if signature(p.text):
            remove_paragraph(p)
        else:
            set_paragraph_text(p, std(p.text))
    for table in list(container.tables):
        clean_table(table)


def clean_docx(path):
    doc = Document(path)
    clean_container(doc)
    for section in doc.sections:
        clean_container(section.header)
        clean_container(section.footer)
    doc.save(path)


def clean_xlsx(path):
    wb = load_workbook(path)
    changed = False
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and not cell.value.startswith('='):
                    nxt = std(cell.value)
                    if nxt != cell.value:
                        cell.value = nxt
                        changed = True
    if changed:
        wb.save(path)
    return changed


def clean_binaries():
    docx_paths = sorted(set(ROOT.glob('cp2025-source-*.docx')) | set((ROOT/'assets/cp-2025/files').glob('*.docx')))
    for path in docx_paths:
        clean_docx(path)
        print('DOCX', path)
    xlsx_paths = sorted(set(ROOT.glob('cp2025-source-*.xlsx')) | set((ROOT/'assets/cp-2025/files').glob('*.xlsx')))
    for path in xlsx_paths:
        if clean_xlsx(path):
            print('XLSX', path)


def clean_text_files():
    text_ext = {'.js','.json','.html','.md','.txt'}
    for path in ROOT.rglob('*'):
        if not path.is_file() or path.suffix.lower() not in text_ext:
            continue
        pstr = path.as_posix()
        if 'assets/quran-kemenag' in pstr or '/.git/' in pstr or pstr.startswith('.git/') or pstr.startswith('.github/'):
            continue
        try:
            old = path.read_text(encoding='utf-8')
        except Exception:
            continue
        new = std(old)
        if path.name.startswith('cp2025-'):
            new = new.replace('CP LAMA 2025 • SUMBER ASLI','CP LAMA 2025 • REFERENSI NASIONAL')
            new = new.replace('Dokumen sumber asli','Dokumen referensi nasional')
            new = new.replace('Buka Berkas Asli','Buka Berkas Referensi')
            new = new.replace('Unduh Dokumen Asli','Unduh Dokumen Referensi')
            new = new.replace('berkas asli','berkas referensi')
            new = new.replace('sumber asli','referensi nasional')
        if new != old:
            path.write_text(new, encoding='utf-8')


def apply_meta(records):
    for rec in records:
        f = ROOT / rec.get('file','')
        if f.is_file():
            data = f.read_bytes()
            rec['size'] = len(data)
            rec['sha256'] = hashlib.sha256(data).hexdigest()


def update_manifests():
    m47 = ROOT/'cp2025-manifest-v47.json'
    if m47.exists():
        data = json.loads(m47.read_text(encoding='utf-8'))
        apply_meta(data.get('records',[]))
        data['source'] = 'CP Lama 2025 — Referensi Nasional'
        m47.write_text(json.dumps(data, ensure_ascii=False, separators=(',',':')), encoding='utf-8')
    m48 = ROOT/'cp2025-manifest-v48.js'
    if m48.exists():
        raw = m48.read_text(encoding='utf-8')
        prefix = 'window.PAIBP_CP2025_V48_MANIFEST = Object.freeze('
        if raw.startswith(prefix) and raw.rstrip().endswith(');'):
            inner = raw[len(prefix):].rstrip()[:-2]
            data = json.loads(inner)
            apply_meta(data.get('records',[]))
            data['source'] = 'CP Lama 2025 — Referensi Nasional'
            m48.write_text(prefix + json.dumps(data, ensure_ascii=False, separators=(',',':')) + ');\n', encoding='utf-8')


def patch_loader():
    app = ROOT/'app-config.js'
    if app.exists():
        t = app.read_text(encoding='utf-8')
        t = t.replace('v81-classroom-access-plus-v80-worship-v79-quran','v82-cp2025-cleanup-plus-v81-classroom')
        t = t.replace('const VERSION = "81";','const VERSION = "82";',1)
        old = 'script("cp2025-loader-v48.js").then(() => script("cp2025-exact-v56.js")).catch(() => false)'
        new = 'script("cp2025-loader-v48.js").then(() => script("cp2025-exact-v56.js")).then(() => script("cp2025-cleanup-v82.js")).catch(() => false)'
        t = t.replace(old,new)
        app.write_text(t,encoding='utf-8')
    for name in ['index.html','akses-guru.html','kendali-editor.html']:
        p = ROOT/name
        if p.exists():
            t = p.read_text(encoding='utf-8')
            t = re.sub(r'app-config\.js\?v=\d+', 'app-config.js?v=82', t)
            p.write_text(t,encoding='utf-8')
    sw = ROOT/'service-worker.js'
    if sw.exists():
        t = sw.read_text(encoding='utf-8')
        t = re.sub(r'const CACHE_NAME="paibp-smart-v\d+[^";]*";', 'const CACHE_NAME="paibp-smart-v82-cp2025-clean-static";', t)
        if 'cp2025-cleanup' not in t:
            t = t.replace('learning-guard|cat-session|service-worker', 'learning-guard|cp2025-cleanup|cp2025-loader|cp2025-exact|cat-session|service-worker')
        sw.write_text(t,encoding='utf-8')


def verify():
    forbidden = ['SMP Negeri 1 Kebonagung','Syaekudin, S.Ag., M.Pd.I.']
    offenders=[]
    for path in ROOT.glob('cp2025-*'):
        if path.is_file() and path.suffix.lower() in {'.js','.json','.html','.md','.txt'}:
            try:
                txt=path.read_text(encoding='utf-8')
            except Exception:
                continue
            for token in forbidden:
                if token.lower() in txt.lower():
                    offenders.append(f'{path}:{token}')
    if offenders:
        raise SystemExit('Forbidden CP2025 text remains: ' + '; '.join(offenders[:20]))


if __name__ == '__main__':
    clean_binaries()
    clean_text_files()
    update_manifests()
    patch_loader()
    verify()
    print('V82 cleanup complete')
